import os
import subprocess
import datetime as dt
import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox

import docx


class InvoiceAutomation:

    def __init__(self):
        self.root = tk.Tk()
        self.root.title("Invoice Automation")
        self.root.geometry("500x600")

        self.client_label = tk.Label(self.root, text="Client Name")
        self.client_street_label = tk.Label(self.root, text="Client Street")
        self.client_suburb_city_zip_label = tk.Label(self.root, text="Client Suburb, City, ZIP")

        self.invoice_number_label = tk.Label(self.root, text="Invoice Number")

        self.service_description_label = tk.Label(self.root, text="Service Description:")
        self.service_hourly_rate_label = tk.Label(self.root, text="Hourly Rate")
        self.service_hours_worked_label = tk.Label(self.root, text="Hours Worked")

        self.payment_method_label = tk.Label(self.root, text="Payment Method:")

        self.payment_methods = {
            'Main Bank': {
                'Recipient': 'The Business Company',
                'Bank': 'ABC Bank',
                'Account Number': '1234567890',
            },
            'Second Bank': {
                'Recipient': 'The Business Company',
                'Bank': 'CBA Bank',
                'Account Number': '1234444444',
            },
            'Third Bank': {
                'Recipient': 'The Business Company',
                'Bank': 'BCA Bank',
                'Account Number': '1233334122',
            }
            
        }

        self.client_entry = tk.Entry(self.root, text="Client Name:")
        self.client_street_entry = tk.Entry(self.root, text="Client Street")
        self.client_suburb_city_zip_entry = tk.Entry(self.root, text="Client Suburb City ZIP")

        self.invoice_number_entry = tk.Entry(self.root, text="Invoice Number")

        self.service_description_entry = tk.Entry(self.root, text="Service Description")
        self.service_hourly_rate_entry = tk.Entry(self.root, text="Hourly Rate")
        self.service_hours_worked_entry = tk.Entry(self.root, text="Service Single Price")

        self.payment_method = tk.StringVar(self.root)
        self.payment_method.set('Main Bank')

        self.payment_method_dropdown = tk.OptionMenu(self.root, self.payment_method, 'Main Bank', "Second Bank", "Third Bank")

        self.create_button = tk.Button(self.root, text="Create Invoice", command=self.create_invoice)

        padding_options= {'fill': 'x', 'expand': True, 'padx': 5, 'pady': 5}

        self.client_label.pack(padding_options)
        self.client_entry.pack(padding_options)
        self.client_street_label.pack(padding_options)
        self.client_street_entry.pack(padding_options)
        self.client_suburb_city_zip_label.pack(padding_options)
        self.client_suburb_city_zip_entry.pack(padding_options)
        self.invoice_number_label.pack(padding_options)
        self.invoice_number_entry.pack(padding_options)
        self.service_description_label.pack(padding_options)
        self.service_description_entry.pack(padding_options)
        self.service_hourly_rate_label.pack(padding_options)
        self.service_hourly_rate_entry.pack(padding_options)
        self.service_hours_worked_label.pack(padding_options)
        self.service_hours_worked_entry.pack(padding_options)

        self.payment_method_label.pack(padding_options)
        self.payment_method_dropdown.pack(padding_options)
        self.create_button.pack(padding_options)

        self.root.mainloop()
    

    @staticmethod
    def replace_text(paragraph, old_text, new_text):
        if old_text in paragraph.text:
            paragraph.text = paragraph.text.replace(old_text, new_text)
    
    
    
    def create_invoice(self):
        doc = docx.Document('invoicetemplate.docx')

        selected_payment_method = self.payment_methods[self.payment_method.get()]

        try:
            replacements = {
                "[Date]": dt.datetime.today().strftime('%Y-%m-%d'),
                "[Client]": self.client_entry.get(),
                "[Client Street]": self.client_street_entry.get(),
                "[Client Suburb City Zip]": self.client_suburb_city_zip_entry.get(),
                "[Invoice Number]": f'#{self.invoice_number_entry.get()}',
                "[Service Description]": self.service_description_entry.get(),
                "[Hourly Rate]": f"${float(self.service_hourly_rate_entry.get()):.2f}",
                "[Hours Worked]": self.service_hours_worked_entry.get(),
                "[Full Price]": f'${float(self.service_hourly_rate_entry.get()) * float(self.service_hours_worked_entry.get()):.2f}',

                "[Recipient]": selected_payment_method['Recipient'],
                "[Bank]": selected_payment_method['Bank'],
                "[Account Number]": selected_payment_method['Account Number'],
                "[Particulars]": self.client_entry.get(),
                "[Reference]": self.invoice_number_entry.get()


            }
        except ValueError:
            messagebox.showerror('Error', 'Invalid amount or price!')
            return
        
        for paragraph in list(doc.paragraphs):
            for old_text, new_text in replacements.items():
                self.replace_text(paragraph, old_text, new_text)
        
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for old_text, new_text in replacements.items():
                            self.replace_text(paragraph, old_text, new_text)
        
        save_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[('PDF documents', '*.pdf')])

        doc.save('filled.docx')

        subprocess.run(['libreoffice', '--headless', '--convert-to', 'pdf', 'filled.docx', '--outdir', '.'], check=True)
        os.rename('filled.pdf', save_path)

        messagebox.showinfo('Success', 'Invoice created and saved successfully!')



if __name__ == "__main__":
    InvoiceAutomation()
    

